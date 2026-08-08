#!/usr/bin/env python3
"""Project a DOCX into repository-native, addressable text files."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document


PROJECTION_VERSION = 1
CHAPTER_RE = re.compile(
    r"^CHAPTER\s+(?P<number>[IVXLCDM]+|\d+)\.?(?:\s*[\u2013\u2014:\-]\s*.+)?$",
    re.IGNORECASE,
)
PAGE_MARKER_RE = re.compile(r"^Page\s+(?P<number>\d+)$", re.IGNORECASE)
GUTENBERG_END_RE = re.compile(r"^\*\*\*\s*End of the project gutenberg", re.IGNORECASE)
WORD_RE = re.compile(r"[^\W_]+(?:[\u2019'\-][^\W_]+)*", re.UNICODE)


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_text(value: str) -> str:
    return sha256_bytes(value.encode("utf-8"))


def normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = value.replace("\u00a0", " ").replace("\u00ad", "")
    return "\n".join(line.rstrip() for line in value.split("\n")).strip()


def roman_to_int(value: str) -> int:
    if value.isdigit():
        return int(value)
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    previous = 0
    for character in reversed(value.upper()):
        current = values[character]
        total += -current if current < previous else current
        previous = max(previous, current)
    return total


def color_value(run: Any) -> str | None:
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def run_record(run: Any) -> dict[str, Any]:
    return {
        "text": normalize_text(run.text),
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
        "font_name": run.font.name,
        "font_size_pt": run.font.size.pt if run.font.size is not None else None,
        "font_color_rgb": color_value(run),
    }


def classify(text: str) -> tuple[str, int | None]:
    if not text:
        return "blank", None
    page = PAGE_MARKER_RE.fullmatch(text)
    if page:
        return "page-marker", int(page.group("number"))
    chapter = CHAPTER_RE.fullmatch(text)
    if chapter:
        return "chapter-heading", roman_to_int(chapter.group("number"))
    if text.upper() in {"PREFACE", "NOTE TO THE THIRD EDITION"}:
        return "front-heading", None
    if GUTENBERG_END_RE.match(text):
        return "gutenberg-end", None
    return "content", None


def safe_slug(value: str) -> str:
    value = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    if not value:
        raise ValueError("Slug must contain an alphanumeric character.")
    return value


def write_json(path: Path, value: Any) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def markdown_for_part(part: dict[str, Any], records: list[dict[str, Any]], source_hash: str) -> str:
    included = [record for record in records if record["kind"] not in {"blank", "page-marker"}]
    lines = [
        "---",
        f'document_sha256: "{source_hash}"',
        f'part: "{part["key"]}"',
        f'source_start: "{part["source_start"]}"',
        f'source_end: "{part["source_end"]}"',
        f"paragraphs: {len(included)}",
        f"words: {sum(record['word_count'] for record in included)}",
        "---",
        "",
    ]

    for record in included:
        lines.append(f'<!-- {record["id"]} -->')
        text = record["text"]
        if record["kind"] in {"chapter-heading", "front-heading"}:
            lines.append(f"# {text}")
        else:
            lines.append(text)
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def project(input_path: Path, output_path: Path, slug: str, source_url: str | None) -> dict[str, Any]:
    input_path = input_path.resolve()
    output_path = output_path.resolve()
    if not input_path.is_file():
        raise FileNotFoundError(input_path)

    source_dir = output_path / "source"
    projection_dir = output_path / "projection"
    chapters_dir = projection_dir / "chapters"
    analysis_dir = output_path / "analysis"
    for directory in (source_dir, chapters_dir, analysis_dir):
        directory.mkdir(parents=True, exist_ok=True)

    source_bytes = input_path.read_bytes()
    source_hash = sha256_bytes(source_bytes)
    source_name = f"{slug}.docx"
    source_copy = source_dir / source_name
    shutil.copyfile(input_path, source_copy)

    document = Document(input_path)
    records: list[dict[str, Any]] = []
    current_part = "front-matter"
    current_chapter: int | None = None
    saw_gutenberg_end = False

    for ordinal, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        kind, marker_number = classify(text)
        if kind == "chapter-heading":
            current_chapter = marker_number
            current_part = f"chapter-{current_chapter:03d}"
        elif kind == "gutenberg-end":
            current_part = "project-gutenberg-tail"
            current_chapter = None
            saw_gutenberg_end = True

        paragraph_id = f"p{ordinal:06d}"
        records.append(
            {
                "id": paragraph_id,
                "ordinal": ordinal,
                "part": current_part,
                "chapter": current_chapter,
                "kind": kind,
                "page_marker": marker_number if kind == "page-marker" else None,
                "style": paragraph.style.name if paragraph.style is not None else None,
                "text": text,
                "sha256": sha256_text(text),
                "character_count": len(text),
                "word_count": len(WORD_RE.findall(text)),
                "runs": [run_record(run) for run in paragraph.runs],
            }
        )

    paragraphs_path = projection_dir / "paragraphs.jsonl"
    with paragraphs_path.open("w", encoding="utf-8", newline="\n") as stream:
        for record in records:
            stream.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")

    table_records: list[dict[str, Any]] = []
    for table_ordinal, table in enumerate(document.tables, start=1):
        rows = []
        for row_ordinal, row in enumerate(table.rows, start=1):
            cells = []
            for cell_ordinal, cell in enumerate(row.cells, start=1):
                cells.append(
                    {
                        "column": cell_ordinal,
                        "text": normalize_text(cell.text),
                        "paragraphs": [normalize_text(paragraph.text) for paragraph in cell.paragraphs],
                    }
                )
            rows.append({"row": row_ordinal, "cells": cells})
        table_records.append(
            {
                "id": f"t{table_ordinal:04d}",
                "ordinal": table_ordinal,
                "style": table.style.name if table.style is not None else None,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "rows": rows,
            }
        )

    tables_path = projection_dir / "tables.jsonl"
    with tables_path.open("w", encoding="utf-8", newline="\n") as stream:
        for record in table_records:
            stream.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")

    grouped: dict[str, list[dict[str, Any]]] = {}
    for record in records:
        grouped.setdefault(record["part"], []).append(record)

    part_entries: list[dict[str, Any]] = []
    for part_key, part_records in grouped.items():
        content_records = [record for record in part_records if record["kind"] not in {"blank", "page-marker"}]
        if not content_records:
            continue
        chapter_number = content_records[0]["chapter"]
        if part_key == "front-matter":
            file_name = "000-front-matter.md"
            label = "Front matter"
        elif part_key == "project-gutenberg-tail":
            file_name = "999-project-gutenberg-tail.md"
            label = "Project Gutenberg tail"
        else:
            file_name = f"{chapter_number:03d}-chapter-{chapter_number:02d}.md"
            label = f"Chapter {chapter_number}"

        entry = {
            "key": part_key,
            "label": label,
            "chapter": chapter_number,
            "file": f"chapters/{file_name}",
            "source_start": part_records[0]["id"],
            "source_end": part_records[-1]["id"],
            "source_paragraph_count": len(part_records),
            "projected_paragraph_count": len(content_records),
            "omitted_blank_count": sum(record["kind"] == "blank" for record in part_records),
            "omitted_page_marker_count": sum(record["kind"] == "page-marker" for record in part_records),
            "word_count": sum(record["word_count"] for record in content_records),
            "character_count": sum(record["character_count"] for record in content_records),
        }
        markdown = markdown_for_part(entry, part_records, source_hash)
        chapter_path = chapters_dir / file_name
        chapter_path.write_text(markdown, encoding="utf-8", newline="\n")
        entry["sha256"] = sha256_text(markdown)
        part_entries.append(entry)

    kind_counts = Counter(record["kind"] for record in records)
    manifest = {
        "projection_version": PROJECTION_VERSION,
        "source": {
            "file": f"source/{source_name}",
            "url": source_url,
            "sha256": source_hash,
            "bytes": len(source_bytes),
        },
        "document": {
            "paragraph_count": len(records),
            "nonempty_paragraph_count": sum(record["kind"] != "blank" for record in records),
            "word_count": sum(record["word_count"] for record in records),
            "character_count": sum(record["character_count"] for record in records),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "chapter_count": sum(entry["chapter"] is not None for entry in part_entries),
            "kinds": dict(sorted(kind_counts.items())),
            "gutenberg_tail_detected": saw_gutenberg_end,
        },
        "ledger": {
            "file": "projection/paragraphs.jsonl",
            "sha256": sha256_bytes(paragraphs_path.read_bytes()),
        },
        "tables": {
            "file": "projection/tables.jsonl",
            "sha256": sha256_bytes(tables_path.read_bytes()),
            "count": len(table_records),
        },
        "parts": part_entries,
    }
    write_json(projection_dir / "manifest.json", manifest)
    write_json(
        source_dir / "source.json",
        {
            "file": source_name,
            "url": source_url,
            "sha256": source_hash,
            "bytes": len(source_bytes),
        },
    )
    return manifest


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path, help="Input DOCX path")
    parser.add_argument("--output", required=True, type=Path, help="Projection root")
    parser.add_argument("--slug", required=True, help="Stable document slug")
    parser.add_argument("--source-url", help="Original download URL")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    manifest = project(args.input, args.output, safe_slug(args.slug), args.source_url)
    summary = {
        "source_sha256": manifest["source"]["sha256"],
        "paragraphs": manifest["document"]["paragraph_count"],
        "words": manifest["document"]["word_count"],
        "chapters": manifest["document"]["chapter_count"],
        "parts": len(manifest["parts"]),
    }
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
